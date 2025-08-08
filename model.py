import os
import fitz  # PyMuPDF
import gc
import logging
import asyncio
import httpx
import base64
from dotenv import load_dotenv
import tempfile
from pathlib import Path
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain.chains import RetrievalQA
from langchain.prompts import ChatPromptTemplate
from langchain.retrievers import BM25Retriever, EnsembleRetriever
from docx import Document as DocxDocument
from urllib.parse import urlparse, unquote
from langchain_core.messages import HumanMessage
import pandas as pd
from io import BytesIO



# Load environment
load_dotenv()
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    raise ValueError("FATAL: OPENAI_API_KEY environment variable not set.")

# Shared resources
EMBEDDINGS_MODEL = OpenAIEmbeddings(model="text-embedding-3-small", openai_api_key=OPENAI_API_KEY)
LLM = ChatOpenAI(model="gpt-4o-mini", openai_api_key=OPENAI_API_KEY)
LLM_VISION = ChatOpenAI(model="gpt-4o-mini", openai_api_key=OPENAI_API_KEY)

FINAL_ANSWER_PROMPT = ChatPromptTemplate.from_messages([
    ("system",
     "You are an AI assistant. Answer questions clearly and concisely in one or two sentences."
     " Base your answers strictly on the provided context. If the answer is not explicitly stated"
     " but can be reasonably inferred, do so carefully. If the context does not support even an inferred answer,"
     " reply with: 'The provided document does not have the answer to the question.'\n\n"
     "Use formal insurance-related language with word and digit formats where relevant.\n"),
    ("human", "Context:\n{context}\n\nQuestion: {question}")
])

REWRITE_PROMPT = ChatPromptTemplate.from_messages([
    ("system", """You are an assistant that rewrites user questions to optimize them for retrieving relevant information from insurance or policy documents.

Rewritten queries must:
- Be concise and keyword-rich
- Eliminate conversational or filler words (e.g., "what", "how", "can I")
- Focus only on the core concepts, terms, or conditions present in formal insurance or policy clauses

Only return the rewritten query without explanations or extra formatting.
"""),
    ("human", "{question}")
])
REWRITE_CHAIN = REWRITE_PROMPT | LLM


def extract_text_from_docx(docx_bytes: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        doc = DocxDocument(tmp.name)
        full_text = "\n".join([para.text for para in doc.paragraphs])
    os.unlink(tmp.name)
    return full_text


def process_pdf_and_create_retriever(pdf_content: bytes) -> EnsembleRetriever:
    with fitz.open(stream=pdf_content, filetype="pdf") as doc:
        full_text = "".join(page.get_text() for page in doc)

    if not full_text.strip():
        raise ValueError("Could not extract text from the PDF.")

    splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=300)
    docs = [Document(page_content=chunk) for chunk in splitter.split_text(full_text)]

    vectordb = FAISS.from_documents(docs, EMBEDDINGS_MODEL)

    dense_retriever = vectordb.as_retriever(search_kwargs={"k": 7})
    bm25_retriever = BM25Retriever.from_documents(docs)
    bm25_retriever.k = 10

    return EnsembleRetriever(retrievers=[dense_retriever, bm25_retriever], weights=[0.7, 0.3])


async def get_final_answer(original_question: str, retriever: EnsembleRetriever) -> str:
    try:
        rewritten_question_response = await REWRITE_CHAIN.ainvoke({"question": original_question})
        rewritten_question = rewritten_question_response.content.strip()
        logging.info(f"Original: '{original_question}' -> Rewritten: '{rewritten_question}'")

        retrieved_docs = await asyncio.to_thread(retriever.get_relevant_documents, rewritten_question)
        context = "\n\n".join(doc.page_content for doc in retrieved_docs)

        final_prompt = FINAL_ANSWER_PROMPT.format_messages(context=context, question=original_question)
        final_response = await LLM.ainvoke(final_prompt)
        return final_response.content.strip()

    except Exception as e:
        logging.exception(f"Error processing question: {original_question}")
        return f"[Error answering question: {str(e)}]"


async def ask_model(file_url: str, questions: list[str]) -> list[str]:
    logging.info(f"Received file URL: {file_url}")
    logging.info(f"Received questions: {questions}")

    try:
        async with httpx.AsyncClient() as client:
            response = await client.get(file_url)
            response.raise_for_status()
            file_bytes = response.content
    except httpx.RequestError as e:
        logging.error(f"Failed to fetch file: {e}")
        return [f"[Error: Could not access the file at {file_url}]"] * len(questions)

    retriever = None
    parsed_url = urlparse(file_url)
    path = unquote(parsed_url.path).lower()

    # --- Handle PDF ---
    if path.endswith(".pdf"):
        try:
            retriever = await asyncio.to_thread(process_pdf_and_create_retriever, file_bytes)
        except Exception as e:
            logging.error(f"Failed to process PDF: {e}")
            return ["[Error: Failed to process the PDF document.]"] * len(questions)

    # --- Handle DOCX ---
    elif path.endswith(".docx"):
        try:
            full_text = await asyncio.to_thread(extract_text_from_docx, file_bytes)
            if not full_text.strip():
                raise ValueError("DOCX file contains no readable text.")

            splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=300)
            docs = [Document(page_content=chunk) for chunk in splitter.split_text(full_text)]

            vectordb = FAISS.from_documents(docs, EMBEDDINGS_MODEL)
            dense_retriever = vectordb.as_retriever(search_kwargs={"k": 7})
            bm25_retriever = BM25Retriever.from_documents(docs)
            bm25_retriever.k = 10
            retriever = EnsembleRetriever(retrievers=[dense_retriever, bm25_retriever], weights=[0.7, 0.3])
        except Exception as e:
            logging.exception("Failed to process DOCX file")
            return ["[Error: Could not process DOCX document.]"] * len(questions)

    # --- Handle Image Files (JPG, PNG, etc.) ---
    elif path.endswith((".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff")):
        logging.info("Processing image file with GPT-4o vision model.")
        try:
            # Convert image to base64
            b64_image = base64.b64encode(file_bytes).decode("utf-8")
            image_url = f"data:image/jpeg;base64,{b64_image}"

            # Construct one vision prompt per question
            tasks = [
                LLM_VISION.ainvoke([
                    HumanMessage(content=[
                        {"type": "image_url", "image_url": {"url": image_url, "detail": "high"}},
                        {"type": "text", "text": question}
                    ])
                ])
                for question in questions
            ]

            responses = await asyncio.gather(*tasks)
            return [res.content.strip() for res in responses]
        except Exception as e:
            logging.exception("Failed during image processing")
            return [f"[Error: Failed to process the image.]"] * len(questions)
        # --- Handle XLSX files ---
    elif path.endswith(".xlsx"):
        logging.info("Processing XLSX file.")
        try:
            # Read Excel from memory
            df_dict = pd.read_excel(BytesIO(file_bytes), sheet_name=None)
            combined_text = ""

            for sheet_name, df in df_dict.items():
                combined_text += f"\n\n--- Sheet: {sheet_name} ---\n"
                combined_text += df.fillna("").astype(str).to_string(index=False)

            if not combined_text.strip():
                raise ValueError("XLSX file contains no readable data.")

            splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=300)
            docs = [Document(page_content=chunk) for chunk in splitter.split_text(combined_text)]

            vectordb = FAISS.from_documents(docs, EMBEDDINGS_MODEL)
            dense_retriever = vectordb.as_retriever(search_kwargs={"k": 7})
            bm25_retriever = BM25Retriever.from_documents(docs)
            bm25_retriever.k = 10
            retriever = EnsembleRetriever(retrievers=[dense_retriever, bm25_retriever], weights=[0.7, 0.3])
        except Exception as e:
            logging.exception("Failed to process XLSX file")
            return ["[Error: Could not process XLSX document.]"] * len(questions)

    else:
        return ["The provided document does not have the answer to the question."] * len(questions)

    # For PDF/DOCX: run RAG QA
    tasks = [get_final_answer(q, retriever) for q in questions]
    answers = await asyncio.gather(*tasks)

    del retriever
    gc.collect()

    return answers